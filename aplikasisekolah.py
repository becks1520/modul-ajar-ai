import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import markdown
from htmldocx import HtmlToDocx

# =====================================
# 1. KONFIGURASI HALAMAN
# =====================================
st.set_page_config(
    page_title="Modul Ajar Berbasis AI",
    page_icon="🎓",
    layout="centered"
)

# =====================================
# 2. STYLE GLOBAL (RESPONSIVE + MOBILE)
# =====================================
st.markdown("""
<style>
.stApp {
    background-image: linear-gradient(to right top, #051937, #004d7a, #008793, #00bf72, #a8eb12);
}
.block-container {
    background: rgba(255,255,255,0.96);
    padding: 26px;
    border-radius: 20px;
    box-shadow: 0 10px 30px rgba(0,0,0,0.25);
    max-width: 1100px;
}
h1 {
    text-align: center;
    font-size: 2.1rem;
    font-weight: 700;
    color: #004d7a;
    margin-bottom: 0.3rem;
}
.form-card {
    background: #ffffff;
    padding: 20px 22px;
    border-radius: 16px;
    box-shadow: 0 6px 16px rgba(0,0,0,0.12);
    margin-bottom: 20px;
}
label {
    font-weight: 600;
    color: #004d7a;
    font-size: 0.9rem;
}
input, textarea, select {
    border-radius: 10px !important;
    padding: 10px 12px !important;
    border: 1px solid #d0d7de !important;
    font-size: 0.95rem !important;
}
input:focus, textarea:focus, select:focus {
    border-color: #008793 !important;
    box-shadow: 0 0 0 2px rgba(0,135,147,0.15) !important;
}
.stButton > button {
    background-color: #004d7a;
    color: white;
    font-weight: bold;
    border-radius: 12px;
    height: 55px;
    font-size: 1rem;
}
@media (max-width: 768px) {
    .block-container { padding: 16px; border-radius: 14px; }
    h1 { font-size: 1.6rem; }
    .form-card { padding: 14px; }
}
</style>
""", unsafe_allow_html=True)

# =====================================
# 3. FUNGSI EXPORT WORD
# =====================================
def create_docx(hasil_ai, data):
    doc = Document()
    header = doc.add_heading(f"MODUL AJAR PRO: {data['mapel']}", 0)
    header.alignment = 1

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Nama Sekolah"
    table.rows[0].cells[1].text = data["sekolah"]
    table.rows[1].cells[0].text = "Mata Pelajaran"
    table.rows[1].cells[1].text = data["mapel"]
    table.rows[2].cells[0].text = "Kelas / Fase"
    table.rows[2].cells[1].text = data["kelas"]
    table.rows[3].cells[0].text = "Materi Pokok"
    table.rows[3].cells[1].text = data["materi"]

    doc.add_paragraph("\n")

    clean_text = hasil_ai.replace("```markdown", "").replace("```", "")
    html_text = markdown.markdown(clean_text, extensions=["tables"])
    HtmlToDocx().add_html_to_document(html_text, doc)

    bio = BytesIO()
    doc.save(bio)
    return bio

# =====================================
# 4. HEADER APLIKASI (BRANDING SEKOLAH)
# =====================================
# PENTING: Jangan ada spasi sebelum tag <div> di bawah ini agar terbaca sebagai desain.
st.markdown("""
<style>
@keyframes fadeUp {
  from { opacity: 0; transform: translateY(18px); }
  to   { opacity: 1; transform: translateY(0); }
}
.school-header {
    text-align: center;
    padding: 3.5rem 1rem 4rem;
    border-radius: 20px;
    background: linear-gradient(135deg, rgba(2,132,199,0.12), rgba(16,185,129,0.12));
    backdrop-filter: blur(12px);
    animation: fadeUp 0.9s ease-out;
    margin-bottom: 1.5rem;
}
.school-badge {
    display: inline-block;
    margin-bottom: 1.2rem;
    padding: 0.35rem 1rem;
    font-size: 0.75rem;
    font-weight: 700;
    letter-spacing: 0.05em;
    color: #0F172A;
    background: rgba(2,132,199,0.22);
    border-radius: 999px;
}
.school-title {
    font-size: 2.8rem;
    font-weight: 800;
    letter-spacing: -0.04em;
    margin-bottom: 0.6rem;
    background: linear-gradient(90deg, #004d7a, #00bf72);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}
.school-subtitle {
    max-width: 820px;
    margin: 0 auto;
    font-size: 1.15rem;
    color: #334155;
    line-height: 1.7;
}
.school-footer {
    margin-top: 1.2rem;
    font-size: 0.85rem;
    color: #475569;
}
@media (max-width: 768px) {
    .school-title { font-size: 2rem; }
    .school-subtitle { font-size: 1rem; }
}
</style>

<div class="school-header">
  <div class="school-badge">✨ Teacher’s Digital Assistant</div>
  <div class="school-title">Modul Ajar AI – SMK Sejahtera</div>
  <div class="school-subtitle">
    Asisten digital guru untuk menyusun Modul Ajar
    Kurikulum Merdeka dengan cara yang lebih mudah,
    cepat, dan praktis.
  </div>
  <div class="school-footer">
    From Teachers, For Teachers
  </div>
""", unsafe_allow_html=True)

st.divider()

# =====================================
# 5. SIDEBAR
# =====================================
with st.sidebar:
    st.header("🔐 Kunci Akses")
    api_key_input = st.text_input("Tempel API Key Gemini", type="password")
    st.info("Mode AI: Auto-Detect")

# =====================================
# 6. FORM INPUT
# =====================================
st.markdown("<div class='form-card'>", unsafe_allow_html=True)
st.subheader("🏫 Identitas Pembelajaran")
col1, col2 = st.columns(2)
with col1:
    sekolah = st.text_input("Nama Sekolah", "SMK Sejahtera")
    kelas = st.text_input("Kelas / Fase", "XI / Fase F")
with col2:
    mapel = st.text_input("Mata Pelajaran", "Produktif / Matematika")
    materi = st.text_input("Materi Pokok", "Materi Inti")
st.markdown("</div>", unsafe_allow_html=True)

st.markdown("<div class='form-card'>", unsafe_allow_html=True)
st.subheader("🧠 Pengaturan Pembelajaran")
col3, col4 = st.columns(2)
with col3:
    model_belajar = st.selectbox("Model Pembelajaran", [
        "Project Based Learning (PjBL)",
        "Problem Based Learning (PBL)",
        "Discovery Learning",
        "Teaching Factory (TeFa)"
    ])
with col4:
    mode_soal = st.selectbox("Jenis Soal", [
        "HOTS - Analisis Kasus",
        "HOTS - Perhitungan",
        "Studi Kasus Industri"
    ])
st.markdown("</div>", unsafe_allow_html=True)

st.markdown("<div class='form-card'>", unsafe_allow_html=True)
st.subheader("👥 Karakteristik Peserta Didik")
murid = st.text_area(
    "Kondisi & Karakteristik Murid",
    placeholder="- 20% mahir\n- 50% sedang\n- 30% perlu bimbingan"
)
st.markdown("</div>", unsafe_allow_html=True)

# =====================================
# 7. GENERATE
# =====================================
if st.button("✨ GENERATE MODUL SUPER LENGKAP", use_container_width=True):

    if not api_key_input or not sekolah or not mapel:
        st.warning("⚠️ API Key, Sekolah, dan Mapel wajib diisi.")
        st.stop()

    genai.configure(api_key=api_key_input)

    found_model = None
    try:
        for m in genai.list_models():
            if "generateContent" in m.supported_generation_methods:
                if "gemini" in m.name.lower():
                    found_model = m.name
                    break
    except Exception as e:
        st.error(f"Gagal mendeteksi model AI: {e}")
        st.stop()

    if not found_model:
        st.error("❌ Tidak ada model Gemini yang mendukung generateContent.")
        st.stop()

    st.toast(f"🤖 Model AI digunakan: {found_model}")
    model = genai.GenerativeModel(found_model)

    prompt = f"""
    Berperan sebagai Konsultan Kurikulum.
    Buatkan Modul Ajar Lengkap Kurikulum Merdeka untuk:
    Sekolah: {sekolah}
    Kelas: {kelas}
    Mapel: {mapel}
    Materi: {materi}
    Model: {model_belajar}
    Kondisi Murid: {murid}
    
    ATURAN FORMATTING (SANGAT PENTING):
    Gunakan format Markdown standar yang rapi (Heading 1, 2, 3, dan Bullet points/Numbering). 
    JANGAN masukkan semua teks ke dalam satu tabel besar. 
    Gunakan tabel HANYA untuk "Kisi-kisi Soal" dan "Rubrik Penilaian" saja.
    JANGAN gunakan format matematika LaTeX (seperti tanda $ atau $$). Tuliskan semua rumus matematika menggunakan teks biasa yang langsung terbaca di Microsoft Word (contoh: x^2 + 5x + 6 = 0).
    
    Sertakan ATP, Soal HOTS, Diferensiasi, dan Rubrik Penilaian.
    """

    with st.spinner("🤖 Menyusun Modul..."):
        hasil = model.generate_content(prompt).text
        st.success("✅ Modul berhasil dibuat!")

        with st.expander("📄 Preview Modul", expanded=True):
            st.markdown(hasil)

        docx = create_docx(hasil, {
            "sekolah": sekolah,
            "mapel": mapel,
            "kelas": kelas,
            "materi": materi
        })

        st.download_button(
            "📥 Download Word (.DOCX)",
            docx.getvalue(),
            file_name=f"Modul_{mapel}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )